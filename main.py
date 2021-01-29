import os
import tempfile
import docx2txt

import click
import docx
from docx import Document

from pathlib import Path
from zipfile import ZipFile

CONTEXT_SETTINGS = dict(help_option_names=['-h', '--help'])

@click.command(context_settings=CONTEXT_SETTINGS)
@click.option('-f','--file', help='.docx file to split')
@click.option('-l','--level', default=1, help='Heading Level to split the file')
def docsplitter(file=None, level=1):
    if(file==None):
        with click.Context(docsplitter) as ctx:
            click.echo(ctx.get_help())
            return
    document = getDocFile(file)
    newBaseName = Path(file).stem
    zipObj = ZipFile(newBaseName + '.zip', 'w')
    with tempfile.TemporaryDirectory() as tmpdirname:
        docx2txt.process(file, tmpdirname)
        rels = {}
        for r in document.part.rels.values():
            if isinstance(r._target, docx.parts.image.ImagePart):
                rels[r.rId] = os.path.basename(r._target.partname)
            i = 0
        currentName = ''
        currentDocument = ''
        for paragraph in document.paragraphs:
            if(is_heading(paragraph, str(level)) or is_Title(paragraph)):
                i+=1
                title = getParagraphTitle(paragraph)
                newName = str(i) + ' - ' + newBaseName + ' ' + title + '.docx'
                if(currentName!='' and currentName != newName):
                    currentDocument.save(os.path.join(tmpdirname, currentName))
                    zipObj.write(os.path.join(tmpdirname, currentName), arcname=currentName)
                newDocument = Document()
                currentName = newName
                currentDocument = newDocument
            cloneParagraph(newDocument, paragraph)
            if 'graphic' in paragraph._p.xml:
                # Get the rId of the image
                for rId in rels:
                    if rId in paragraph._p.xml:
                        newDocument.add_picture(os.path.join(tmpdirname, rels[rId]))
    zipObj.close()

def is_heading(paragraph, level):
    if paragraph.style.name.startswith('Heading '+str(level)):
        return True
    return False

def is_Title(paragraph):
    if paragraph.style.name.startswith('Title'):
        return True
    return False

def getParagraphTitle(paragraph):
        return paragraph.text

def getDocFile(filename):
    return Document(filename)

def cloneParagraph(document, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = document.add_paragraph(paragraph.text)
    for run in paragraph.runs:
        # output_run = output_para.add_run(run.text)
        output_run = output_para.add_run()

        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.font.name = run.font.name
        output_run.font.highlight_color = run.font.highlight_color
        output_run.font.size = run.font.size
        # output_run.style.name = run.style.name
        output_run.style = run.style

    # Paragraph's alignment data
    output_para.style = paragraph.style
    output_para.alignment = paragraph.alignment
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control

if __name__ == '__main__':
    docsplitter()