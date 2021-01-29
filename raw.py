import glob

import docx
from docx import Document
from docx.document import Document as _Document
import xml.etree.ElementTree as ET
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.oxml import CT_P, CT_Tbl
from docx.text.paragraph import Paragraph
import docx2txt
import hashlib
import tempfile
import os
from zipfile import ZipFile

# create a ZipFile object
zipObj = ZipFile('sample.zip', 'w')
# Add multiple files to the zip
zipObj.write('sample_file.csv')
zipObj.write('test_1.log')
zipObj.write('test_2.log')
# close the Zip File
zipObj.close()

def hash_file(filename):
   """"This function returns the SHA-1 hash
   of the file passed into it"""

   # make a hash object
   h = hashlib.sha1()

   # open file for reading in binary mode
   with open(filename,'rb') as file:

       # loop till the end of the file
       chunk = 0
       while chunk != b'':
           # read only 1024 bytes at a time
           chunk = file.read(1024)
           h.update(chunk)

   # return the hex representation of digest
   return h.hexdigest()

# file_hash = hash_file('edebiyat.docx');

with tempfile.TemporaryDirectory() as tmpdirname:


if not os.path.exists('./images'):
    os.makedirs('./images')

docx2txt.process('edebiyat.docx', './images' )

document = Document('edebiyat.docx')
document_part = document.inline_shapes.part

rels = {}
for r in document.part.rels.values():
    if isinstance(r._target, docx.parts.image.ImagePart):
        rels[r.rId] = os.path.basename(r._target.partname)
#print(temp_dir.name)

# Then process your text
for paragraph in document.paragraphs:
    # If you find an image
    if 'graphic' in paragraph._p.xml:
        # Get the rId of the image
        for rId in rels:
            if rId in paragraph._p.xml:
                print(os.path.join('./images', rels[rId]))
    else:
        print(paragraph.text)

#print(glob.glob(temp_dir.name + "/*"))

# use temp_dir, and when done:
#temp_dir.cleanup()
#print(type(document._body._body))

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def hasImage(par):
    """get all of the images in a paragraph
    :param par: a paragraph object from docx
    :return: a list of r:embed
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
        ids.append(id)

    return ids

def getText(document):
    fullText = []
    for para in document.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getSectionText(section):
    fullText = []
    for para in section:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getSectionTitle(section):
    return section[0].text

def iterate_document_sections(document):
    """Generate a sequence of paragraphs for each headed section in document.

    Each generated sequence has a heading paragraph in its first position,
    followed by one or more body paragraphs.
    """
    paragraphs = [document.paragraphs[0]]
    for paragraph in document.paragraphs[1:]:
        if is_heading(paragraph):
             yield paragraphs
             paragraphs = [paragraph]
             continue
        paragraphs.append(paragraph)
    yield paragraphs

def add_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        # output_run = output_para.add_run(run.text)
        output_run = output_para.add_run()

        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.font.name = run.font.name
        output_run.font.highlight_color = run.font.highlight_color
        output_run.font.size = run.font.size
        # output_run.style.name = run.style.name
        output_run.style = run.style



    # Paragraph's alignment data
    output_para.style = paragraph.style
    output_para.alignment = paragraph.alignment
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control

def is_heading(paragraph):
    if paragraph.style.name.startswith('Heading 1'):
        return True
    return False



def create_document_from_paragraphs(paragraphs):
    """Iterates through the paragraphs containing articles and splits them into separate files.

    :param paragraphs: Article text
    :return: New document with a single article.
    """
    new_doc = Document()
    for counter, words in enumerate(paragraphs):
        new_content = words.text
        new_doc.add_paragraph(new_content)
        print (new_content)
    new_doc.save('Articles/new_doc' + str(counter) + '.docx')

"""
for element in sub_doc._document_part.body._element:
    combined_document._document_part.body._element.append(element)

for element in sub_doc.element.body:
                merged_document.element.body.append(element)
"""